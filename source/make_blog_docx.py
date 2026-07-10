# -*- coding: utf-8 -*-
"""
무이성형외과 유튜브 → 네이버 블로그 포스팅 초안 생성기
소재 영상: https://www.youtube.com/watch?v=-aP9Jo7MmAA  (무턱 / 턱끝 전진)
어투·양식 기준: 기존 무이성형외과 블로그 포스팅 (변형 T절골 편)
출력: ../무이성형외과_블로그초안_무턱_턱끝전진.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

FONT = "Malgun Gothic"
INK = RGBColor(0x2B, 0x26, 0x22)     # 브랜드 다크
ACCENT = RGBColor(0xB0, 0x8A, 0x5E)  # 포인트(모카)
MUT = RGBColor(0x8A, 0x81, 0x78)     # 안내 회색
BODY = RGBColor(0x33, 0x30, 0x2C)

VIDEO_URL = "https://www.youtube.com/watch?v=-aP9Jo7MmAA"


def _kfont(run, size=11, bold=False, color=BODY, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), FONT)
    return run


def para(doc, text="", size=11, bold=False, color=BODY, align=None,
         before=4, after=8, italic=False, line=1.5):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if text:
        _kfont(p.add_run(text), size=size, bold=bold, color=color, italic=italic)
    return p


def heading(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(20)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.4
    _kfont(p.add_run(text), size=14, bold=True, color=INK)
    return p


def quote(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.4
    pf.left_indent = Cm(0.4)
    _kfont(p.add_run(text), size=11.5, bold=True, color=INK)
    return p


def imgph(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(10)
    _kfont(p.add_run(f"[ {text} ]"), size=10.5, color=MUT, italic=True)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.4
    _kfont(p.add_run(text), size=11, color=BODY)
    return p


def faq(doc, q, a):
    pq = doc.add_paragraph()
    pq.paragraph_format.space_before = Pt(10)
    pq.paragraph_format.space_after = Pt(2)
    pq.paragraph_format.line_spacing = 1.4
    _kfont(pq.add_run("Q. "), size=11, bold=True, color=ACCENT)
    _kfont(pq.add_run(q), size=11, bold=True, color=INK)
    pa = doc.add_paragraph()
    pa.paragraph_format.space_after = Pt(4)
    pa.paragraph_format.line_spacing = 1.5
    _kfont(pa.add_run("A. "), size=11, bold=True, color=ACCENT)
    _kfont(pa.add_run(a), size=11, color=BODY)


def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pbdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pbdr.makeelement(qn('w:bottom'), {})
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'D9D2C8')
    pbdr.append(bottom); pPr.append(pbdr)


# ─────────────────────────────────────────────
doc = Document()
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(11)
normal.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

# 상단 안내(초안 메타)
para(doc, "무이성형외과 · 유튜브 영상 기반 네이버 블로그 포스팅 초안",
     size=10.5, color=MUT, after=2, before=0)
para(doc, f"소재 영상: {VIDEO_URL}   |   작성일: 2026-07-10",
     size=10.5, color=MUT, after=2, before=0)
para(doc, "※ [ ] 표시는 이미지 삽입 위치입니다. 게시 전 의료광고 규정에 따른 최종 검수를 권장드립니다.",
     size=10.5, color=MUT, after=6, before=0)
divider(doc)

# 제목
para(doc, "무턱, 보형물 없이 개선할 수 있을까? – 턱끝 전진으로 인상까지 살리는 윤곽수술",
     size=17, bold=True, color=INK, before=2, after=14, line=1.3)

# 도입 인용
quote(doc, "“무턱이라 그런지, 마스크를 벗으면 유독 옆모습에 자신이 없어요.”")
quote(doc, "“턱끝이 짧고 뒤로 밀려 있어서, 인상이 약해 보인다는 말을 종종 들어요.”")
quote(doc, "“필러나 보형물 말고, 제 뼈로 자연스럽게 바꿀 방법은 없을까요?”")

para(doc, "무턱 상담에서는 단순히 ‘턱이 짧다’는 이야기만 하고 끝나는 분이 드뭅니다. "
          "턱끝이 뒤로 물러나 있어 옆모습이 밋밋하고, 그래서 인상이 다소 옹졸해 보이며, "
          "자꾸 마스크를 찾게 된다고 말씀하시는 분들이 많거든요.", before=10)

para(doc, "결론부터 말씀드리면, 무턱은 무조건 보형물로 채워 넣는 게 아니라 "
          "내 턱뼈를 앞으로 전진시켜 길이감과 입체감을 만들 때 자연스럽게 개선되는 윤곽수술에 가깝습니다.")

para(doc, "오늘은 유일무이한 아름다움, 무이성형외과✨에서 무턱을 어떤 기준으로 바라보고, "
          "턱끝 전진으로 인상을 어떻게 살리는지 솔직하게 풀어드릴게요.")
imgph(doc, "썸네일 이미지 삽입")

# 1
heading(doc, "무턱, 왜 인상까지 달라 보일까?")
para(doc, "특히 남성분들은 무턱이면 인상이 다소 옹졸해 보이고, 남성적인 느낌이 떨어져 보이는 경우가 많습니다. "
          "그러다 보니 마스크를 자주 끼게 되고, 어딘가 자신감이 줄어든다고 느끼시는 분들도 계세요.")
para(doc, "턱끝이 충분히 자리 잡지 못하면 얼굴 아래쪽의 무게 중심이 흐려져, 전체적인 인상이 약해 보이기 쉽습니다. "
          "그래서 무턱은 단순히 ‘턱이 짧은 문제’라기보다, 얼굴 전체 인상과 연결된 부분으로 봐야 해요.")
imgph(doc, "유튜브 영상 캡쳐 이미지 삽입")

# 2
heading(doc, "무턱 개선의 핵심, 턱끝을 ‘앞으로’ (전진)")
para(doc, "턱끝은 앞으로 전진시킬 수도, 뒤로 후진시킬 수도 있고, 길이감을 늘리거나 줄이는 등 방법이 굉장히 다양합니다. "
          "그중 무턱은 턱끝을 조금 앞쪽으로 전진시켜 개선하는 방향이 핵심이에요.")
para(doc, "뼈에서는 턱끝을 앞쪽으로 전진시킨 뒤 단단히 고정하게 되는데, 이렇게 하면 부족했던 턱의 볼륨감과 입체감이 살아나게 됩니다. "
          "없던 라인을 억지로 만드는 게 아니라, 뒤로 물러나 있던 내 턱뼈의 위치를 앞으로 옮겨 준다고 이해하시면 쉬워요.")
imgph(doc, "유튜브 영상 이미지 삽입 – 턱끝 전진 전/후")

# 3
heading(doc, "얼마나 나와야 자연스러울까? – 에스테틱 라인")
para(doc, "턱끝이 어느 정도 나와야 조화로운지는 ‘에스테틱 라인(Aesthetic line)’을 기준으로 봅니다. "
          "무턱은 이 라인에 턱끝이 못 미치는 경우가 많은데, 수술은 에스테틱 라인에 맞춰 턱을 전진시키는 방향으로 설계돼요.")
para(doc, "남성분의 경우 보통 짧게는 6mm, 길게는 8mm 정도까지 전진을 고려할 수 있습니다. "
          "다만 이 수치는 얼굴 비율과 뼈 구조에 따라 달라지기 때문에, 정해진 정답이 아니라 개인에 맞춰 조정되는 값이에요.")

# 4
heading(doc, "남성만? 여성 무턱도 함께 볼 수 있어요")
para(doc, "무턱은 남성분만의 고민이 아닙니다. 여성분들도 무턱이 심하면 턱 아래가 접히듯 보여 "
          "‘이중턱 같다’고 느끼시는 경우가 있어요.")
para(doc, "이런 경우에도 턱끝을 전진시켜 입체감 있는 라인을 만들면, 밋밋하던 하관에 자연스러운 볼륨이 생기면서 "
          "옆모습 인상이 한결 정돈될 수 있습니다.")

# 5
heading(doc, "턱끝 필러·보형물 없이, 내 뼈로")
para(doc, "무턱이라고 하면 흔히 턱끝 필러나 턱끝 보형물부터 떠올리시는 경우가 많습니다. "
          "하지만 이렇게 뼈 자체를 앞으로 전진시켜 고정하는 방법으로도 무턱을 보완할 수 있어요.")
para(doc, "이물질을 넣어 볼륨을 채우는 방식과 달리, 내 뼈의 위치를 옮겨 길이감과 입체감을 함께 만든다는 점에서 접근이 다릅니다. "
          "그래서 필러나 보형물이 부담스러웠던 분들도 한 번쯤 상담을 통해 확인해 보실 수 있는 방법이에요.")
imgph(doc, "유튜브 영상 이미지 삽입")

# 6
heading(doc, "이런 분들에게 잘 맞을 수 있어요")
bullet(doc, "옆모습에서 턱끝이 뒤로 밀려 있어 밋밋해 보이는 분")
bullet(doc, "무턱 때문에 인상이 다소 옹졸해 보인다고 느끼는 남성분")
bullet(doc, "무턱이 심해 이중턱처럼 보이는 여성분")
bullet(doc, "필러·보형물 대신 내 뼈로 무턱을 개선하고 싶은 분")
bullet(doc, "과한 변화보다 자연스러운 입체감과 인상 변화를 원하는 분")

# 7
heading(doc, "다만, 알아두셔야 할 점도 있어요")
para(doc, "턱끝 전진술은 뼈를 다루는 윤곽수술인 만큼, 수술 전 정확한 CT·X-ray 진단과 신경 구조 확인이 반드시 선행되어야 합니다. "
          "뼈의 두께와 신경선 위치, 얼굴 전체 비율에 따라 전진할 수 있는 범위와 방법이 개인마다 달라질 수 있어요.")
para(doc, "그래서 중요한 건 ‘무조건 많이 넣거나 많이 빼는 것’이 아니라, "
          "지금 내 뼈 구조에 맞는 안전한 계획을 세우는 것입니다.")

# 8 FAQ
heading(doc, "자주 묻는 질문 (FAQ)")
faq(doc, "무턱인데 꼭 보형물(실리콘)을 넣어야 하나요?",
    "그렇지 않아요. 턱끝 뼈를 앞으로 전진시켜 고정하는 방법으로도 무턱을 개선할 수 있는 경우가 많습니다. "
    "다만 뼈 구조에 따라 더 적합한 방법은 달라질 수 있어, 진단 후 결정하는 것이 좋아요.")
faq(doc, "턱은 얼마나 앞으로 나올 수 있나요?",
    "남성분 기준으로 보통 짧게는 6mm, 길게는 8mm 정도까지 전진을 고려할 수 있습니다. "
    "개인의 뼈·얼굴 비율에 따라 범위는 달라질 수 있어요.")
faq(doc, "무턱 수술은 남자만 하는 건가요?",
    "아니에요. 여성분도 무턱이 심해 이중턱처럼 보이는 경우, 턱끝 전진으로 입체감 있는 라인을 만들어 볼 수 있습니다.")
faq(doc, "상담 때 꼭 필요한 게 있을까요?",
    "CT·X-ray 촬영이 핵심이에요. 신경선 위치와 뼈 구조를 확인해야 안전한 수술 범위를 정할 수 있기 때문에, "
    "정확한 진단부터 시작하시길 권해드려요.")

# 정리
heading(doc, "정리하면,")
para(doc, "자연스러운 무턱 개선은 ‘무조건 채워 넣는 것’보다, 내 뼈를 에스테틱 라인에 맞춰 앞으로 전진시키는 "
          "‘설계’에서 시작됩니다.")
para(doc, "무턱 때문에 옆모습이 아쉬우시거나 인상이 약해 보여 고민이시라면, 혼자 판단하지 마시고 "
          "정확한 CT 검사로 나에게 맞는 안전한 계획을 세워 보시길 바랍니다.")
para(doc, "무이성형외과에서는 정면뿐 아니라 측면·하안면 비율과 신경선 위치까지 함께 확인한 뒤, "
          "안전은 지키고 인상과 입체감은 살리는 방향으로 턱끝 전진술을 디자인합니다.")

# 영상 링크
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after = Pt(10)
_kfont(p.add_run(f"▶ 실제 수술 원리는 영상에서 확인하실 수 있어요 : {VIDEO_URL}"),
       size=11, bold=True, color=INK)

# 해시태그
para(doc, "#무턱 #무턱수술 #턱끝전진 #턱끝성형 #남성무턱 #무턱보형물 #이중턱 #에스테틱라인 "
          "#윤곽수술 #V라인윤곽수술 #강남윤곽수술 #강남턱수술 #강남성형외과 #무이성형외과 #무이윤곽수술",
     size=10.5, color=ACCENT, before=6, after=14, line=1.6)

divider(doc)
# 고지 문구
para(doc, "본 내용은 시술 및 수술 관련 정보를 이해하기 쉽게 전달하기 위한 정보 전달 목적으로 "
          "의료법 제56조 1항을 준수하여, 무이성형외과의원에서 직접 작성했습니다. "
          "모든 시/수술은 개인마다 차이 및 부작용이 발생할 수 있으니 반드시 담당 의료진과 상담 후 진행하시길 권장드립니다.",
     size=9.5, color=MUT, line=1.5)

out = os.path.join(os.path.dirname(__file__), "..",
                   "무이성형외과_블로그초안_무턱_턱끝전진.docx")
out = os.path.abspath(out)
doc.save(out)
print("saved:", out)
