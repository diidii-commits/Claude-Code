# -*- coding: utf-8 -*-
"""
무이성형외과 블로그 초안 — 나이대별 성형 (성형학개론 영상 기반) v3
문체 기준: 224333520810 (인용 3개 → 공감 → 결론부터 → 시그니처 → 대화체 소제목 → FAQ → 정리하면)
출력: ../무이성형외과_블로그초안_나이대별성형_리라이트.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn

FONT = "Malgun Gothic"
INK = RGBColor(0x2B, 0x26, 0x22)
GREEN = RGBColor(0x24, 0x5B, 0x12)
ACCENT = RGBColor(0xB0, 0x8A, 0x5E)
MUT = RGBColor(0x8A, 0x81, 0x78)
BODY = RGBColor(0x33, 0x30, 0x2C)
CAP = RGBColor(0x9A, 0x7A, 0x4E)

VIDEO_URL = "https://www.youtube.com/watch?v=zcsPkIxab_8"


def _kf(run, size=11, bold=False, color=BODY, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = rpr.makeelement(qn('w:rFonts'), {})
        rpr.append(rf)
    rf.set(qn('w:eastAsia'), FONT)
    return run


def para(doc, text="", size=11, bold=False, color=BODY, before=3, after=7, italic=False, line=1.5):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = line
    if text:
        _kf(p.add_run(text), size=size, bold=bold, color=color, italic=italic)
    return p


def heading(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(20); pf.space_after = Pt(6); pf.line_spacing = 1.4
    _kf(p.add_run(text), size=14, bold=True, color=GREEN)
    return p


def quote(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(2); pf.space_after = Pt(2); pf.line_spacing = 1.4
    pf.left_indent = Cm(0.4)
    _kf(p.add_run(text), size=11.5, bold=True, color=INK)
    return p


def capture(doc, ts, scene, alt):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10); pf.space_after = Pt(2)
    _kf(p.add_run(f"[영상 캡처 — {ts}  {scene}]"), size=10, bold=True, color=CAP)
    pa = doc.add_paragraph()
    pa.paragraph_format.space_after = Pt(9)
    _kf(pa.add_run(f"ALT: {alt}"), size=9.5, color=MUT, italic=True)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.space_before = Pt(1); pf.space_after = Pt(3); pf.line_spacing = 1.4
    _kf(p.add_run(text), size=11, color=BODY)


def faq(doc, q, a):
    pq = doc.add_paragraph()
    pq.paragraph_format.space_before = Pt(10); pq.paragraph_format.space_after = Pt(2)
    pq.paragraph_format.line_spacing = 1.4
    _kf(pq.add_run("Q. "), size=11, bold=True, color=ACCENT)
    _kf(pq.add_run(q), size=11, bold=True, color=INK)
    pa = doc.add_paragraph()
    pa.paragraph_format.space_after = Pt(4); pa.paragraph_format.line_spacing = 1.5
    _kf(pa.add_run("A. "), size=11, bold=True, color=ACCENT)
    _kf(pa.add_run(a), size=11, color=BODY)


def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pbdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pbdr.makeelement(qn('w:bottom'), {})
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'D9D2C8')
    pbdr.append(bottom); pPr.append(pbdr)


def table(doc, rows, widths, header=True):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row):
            c = t.rows[ri].cells[ci]
            c.width = Cm(widths[ci])
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
            _kf(p.add_run(cell), size=10,
                bold=(header and ri == 0) or ci == 0,
                color=GREEN if (header and ri == 0) else BODY)
    return t


doc = Document()
normal = doc.styles["Normal"]
normal.font.name = FONT; normal.font.size = Pt(11)
normal.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

# ── 상단 메타
para(doc, "무이성형외과 · 「성형학개론」 영상 기반 — 나이대별 성형 (리라이트 v3, 문체 교정판)", size=10.5, color=MUT, after=2, before=0)
para(doc, f"소재 영상: {VIDEO_URL}   |   2026-08-05   |   발행 전 검수용", size=10.5, color=MUT, after=6, before=0)
divider(doc)

# ── 제목
para(doc, "제목 (추천)", size=12, bold=True, color=ACCENT, after=2)
para(doc, "나이대별 성형, 20대 30대 40대 나에게 맞는 순서 찾는 법", size=16, bold=True, color=INK, after=2, line=1.3)
para(doc, "30자 내(검색결과 미절단). '나에게 맞는 ~ 찾는 법'은 이 블로그 1,138회 글의 검증 패턴. "
          "예비: 「20대 30대 40대 성형, 나이대별로 순서 정하는 법」", size=9.5, color=MUT, after=10)
divider(doc)

# ── 본문
para(doc, "본문 (붙여넣기용)", size=12, bold=True, color=ACCENT, after=6)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(9)
_kf(p.add_run("[영상 임베드 — 글 상단]  스마트에디터 → 동영상 → 아래 URL 붙여넣기 → "), size=10, bold=True, color=CAP)
_kf(p.add_run(VIDEO_URL), size=10, color=CAP)

quote(doc, "“지금 뭘 해야 할지는 모르겠는데, 뭔가는 해야 할 것 같아요.”")
quote(doc, "“20대에 눈은 했는데, 30대엔 뭘 봐야 하나요?”")
quote(doc, "“이 나이에 시작하면 너무 늦은 걸까요?”")

para(doc, "상담실에서 정말 자주 듣는 이야기입니다. 가만히 들어보면 부위가 고민이 아니라, "
          "'순서'가 고민이신 분들이 생각보다 많거든요.", before=9)
para(doc, "결론부터 말씀드리면, 나이대별 성형에 정해진 정답표는 없습니다. "
          "다만 나이대마다 고민의 중심이 옮겨가는 흐름은 분명히 있어서, "
          "이 지도를 알고 나면 지금 내 나이에서 뭘 먼저 봐야 할지 훨씬 선명해져요.", bold=True)
para(doc, "오늘은 유일무이한 아름다움, 무이성형외과✨에서 "
          "10대부터 50대까지 나이대별 성형 고민이 어떻게 달라지는지, 솔직하게 풀어드릴게요.")

capture(doc, "0:15", "화이트보드에 남녀 나이대별 표를 그려놓고 설명을 시작하는 장면",
        "나이대별 성형 상담 관심사를 성별과 연령대로 정리한 표를 설명하는 장면")

heading(doc, "한눈에 보는 나이대별 성형 지도")
para(doc, "먼저 전체 지도부터 펼쳐볼게요. 아래 표는 상담실에서 보이는 '관심의 중심'을 정리한 것이지, "
          "권장 순서표가 아니라는 점만 기억해 주세요. 사람마다 순서는 얼마든지 달라질 수 있으니까요.", after=6)
table(doc, [
    ["나이대", "여성 — 관심의 중심", "남성 — 관심의 중심"],
    ["10대", "눈 (쌍꺼풀), 일부는 코와 얼굴형까지", "아직 관심이 적은 시기"],
    ["20대", "눈 구체화, 사각턱 보톡스", "연애를 계기로 시작 — 무쌍, 눈매교정, 코"],
    ["20대 후반~30대", "윤곽 첫 수술 + 쌍꺼풀 재수술, 시술 병행", "상대적으로 조용한 시기"],
    ["40~50대", "상안검, 하안검, 거상 — 시술 한계를 느끼는 시점", "50대에 상안검, 하안검, 거상"],
], widths=[3.2, 6.6, 6.2])
para(doc, "※ 스마트에디터의 표 기능으로 다시 만들어 넣어주세요. 표가 있으면 읽는 분들 체류시간에도 도움이 됩니다.",
     size=9.5, color=MUT, before=4)

heading(doc, "10대 20대 성형, 왜 다들 눈부터 시작할까요?")
para(doc, "여성분들은 사실 청소년기부터 고민이 시작됩니다. 친구들끼리 거울을 보면서 "
          "\"너는 쌍꺼풀 있는 게 어울린다\", \"너는 여기만 하면 되겠다\" 같은 얘기, 다들 한 번쯤 해보셨을 거예요. "
          "그래서 10대는 눈에 대한 관심이 가장 큽니다.")
para(doc, "한 발 더 나간 친구들은 코까지 봐요. 면봉으로 콧대를 눌러보면서 서로 봐주기도 하고요. "
          "일부는 이때부터 사각턱이나 광대 같은 얼굴형 고민을 시작합니다.")
capture(doc, "1:07", "\"근데 이때부터 사실은 알아요\"라고 말하는 장면",
        "10대부터 자기 얼굴의 고민 부위를 이미 알고 있다는 나이대별 성형 상담 이야기")
para(doc, "재밌는 건, 이 시기에 이미 본인이 답을 알고 있다는 거예요. 내가 눈이 문제인지, 코인지, 광대나 사각턱인지. "
          "그래서 상담은 새로 발견하는 자리라기보다 확인받는 자리에 가깝습니다. "
          "혹시 지금 마음속에 짚이는 부위가 있으시다면, 아마 그게 정답일 가능성이 높아요.")
para(doc, "20대가 되면 큰 틀은 그대로인데 결이 달라집니다. 화장이라는 무기가 생기거든요. "
          "셰딩으로 얼굴형을 보완하고, 아이라인으로 눈매를 잡고요. 그래서 원하는 것도 훨씬 구체적으로 바뀝니다. "
          "\"아이라인이 지워지니까 쌍꺼풀을 조금 키우고 싶어요\", \"눈밑을 정돈하고 싶어요\" 같은 식으로요.")

heading(doc, "20대 후반 30대 성형, 조합이 달라지는 시기예요")
para(doc, "얼굴형 쪽은 사각턱 보톡스를 이미 맞아보신 분들이 많습니다. 보톡스로 한계가 느껴지면 그때부터 윤곽 성형을 고민하기 시작하고, "
          "사각턱과 함께 광대도 같이 보게 되죠. 실제로 윤곽 상담을 가장 많이 오시는 시기가 20대 후반에서 30대예요.")
para(doc, "그리고 이때 새로 등장하는 게 쌍꺼풀 재수술입니다. 비절개로 잡았던 라인이 낮아지기도 하고, "
          "절개로 했는데 모양이 아쉬운 경우도 있거든요.")
capture(doc, "2:41", "\"윤곽은 처음인데 눈은 재수술\"이라고 설명하는 장면",
        "30대 나이대별 성형 상담에서 윤곽은 첫 수술, 눈은 재수술로 겹치는 패턴")
para(doc, "그래서 이 시기 상담은 조합이 독특해요. 윤곽은 처음인데, 눈은 재수술. 이런 분들이 정말 많습니다. "
          "경제적으로 여유가 생기면서 울쎄라나 써마지 같은 시술 정보도 이미 알고 오시는 경우가 많고요. "
          "수술과 시술을 함께 저울질하기 시작하는 시기라고 보시면 됩니다.")

heading(doc, "40대 성형 상담, 눈꺼풀과 처짐이 중심이 됩니다")
para(doc, "40대가 되면 상안검, 하안검이 상담 테이블에 올라옵니다. 시작이 빠른 분들은 거상까지 함께 보시고요.")
capture(doc, "3:22", "울쎄라·써마지 샷수 이야기를 하는 장면",
        "울쎄라 써마지 반복 후 한계를 느끼고 40대 상안검 거상을 고민하게 되는 흐름")
para(doc, "계기는 대체로 비슷합니다. 울쎄라나 써마지를 300샷 정도 받을 때는 변화가 느껴졌는데, "
          "몇 년 지나 600샷, 800샷을 해도 예전 같지 않다고 느껴지는 순간이 오거든요. "
          "그때 흉터 걱정과 함께 수술적인 방법을 처음 고민하게 됩니다.")
para(doc, "혹시 '시술을 해도 예전만큼 달라지지 않네' 싶으시다면, 그게 하나의 신호일 수 있어요. "
          "여기서부터는 한 번에 끝내는 게 아니라, 관리라는 관점에서 길게 보셔야 하는 시기입니다.")

heading(doc, "남자 나이대별 성형, 흐름이 꽤 다릅니다")
para(doc, "남성분들은 시작 자체가 늦어요. 성형이 처음 등장하는 시점은 보통 20대, 계기는 대체로 연애입니다. "
          "본인 기준으로는 괜찮다고 생각했는데 객관적인 평가를 마주하게 되면서, 그때부터 친구들과 얼굴 이야기를 시작하거든요.")
capture(doc, "5:16", "남성이 고민하는 부위를 하나씩 짚는 장면",
        "남자 나이대별 성형에서 무쌍 눈매교정 코성형 사각턱을 고민하는 20대 패턴")
para(doc, "이때 나오는 이야기가 무쌍 성형, 눈매교정, 코, 그리고 얼굴 크기 정리예요. "
          "특징이 하나 있다면 주변 피드백에 굉장히 민감하다는 점입니다. 여성분들은 원하는 방향을 밀고 가시는 편인데, "
          "남성분들은 \"그렇게 하면 별로다\" 한마디에 방향을 바꾸시는 경우가 많거든요.")
para(doc, "30대에는 오히려 조용해집니다. 결혼을 바라보는 시기가 되면서 외모 고민의 비중이 줄어들어요. "
          "이 시기 남성 상담은 배우자나 여자친구를 따라왔다가 \"제 눈밑도 좀 어두운가요?\" 하고 물어보시는 형태가 대부분입니다.")
capture(doc, "6:56", "남성 상담의 결정 속도를 이야기하는 장면",
        "50대 남자 상안검 하안검 거상 상담에서 나타나는 결정 패턴")
para(doc, "그리고 50대에 다시 등장합니다. 상안검, 하안검, 거상이 함께 언급되는데, 배우자분과 함께 오셔서 상담받는 형태가 흔해요. "
          "결정도 빠른 편입니다. 방향이 잡히면 오래 고민하지 않고 진행하시고, 회복 과정을 감수하면서 관리하시는 남성분들도 요즘은 꽤 늘었습니다.")

heading(doc, "이런 분들이라면 상담을 권해드려요")
bullet(doc, "뭔가 하고 싶은데 어디부터 봐야 할지 순서가 잡히지 않는 분")
bullet(doc, "20대 후반 이후 눈 재수술과 윤곽을 동시에 고민 중인 분")
bullet(doc, "시술을 반복해도 예전만큼 변화가 느껴지지 않는 분")
bullet(doc, "40대 이후 상안검, 하안검 시기를 언제로 잡을지 궁금하신 분")
bullet(doc, "가족이나 배우자와 함께 상담받고 싶으신 분")

heading(doc, "다만, 알아두셔야 할 점도 있어요")
para(doc, "나이대별 흐름은 어디까지나 경향이에요. 개인의 뼈 구조와 피부 상태, 이전 수술 이력에 따라 "
          "필요한 방법과 시기는 완전히 달라집니다. 특히 재수술이 포함되거나 눈꺼풀, 처짐이 얽힌 경우에는 "
          "진찰과 검사 없이 순서를 정하기 어려워요. 그래서 순서를 정하는 일도 결국, 정확한 진단이 먼저입니다.")

heading(doc, "자주 묻는 질문 (FAQ)")
faq(doc, "나이대별 성형에 정해진 순서가 있나요?",
    "정해진 표는 없어요. 다만 상담 현장에서는 눈, 윤곽과 재수술, 눈꺼풀과 처짐 순으로 관심이 옮겨가는 경향이 있습니다. "
    "순서는 지금 본인이 가장 불편하게 느끼는 부위를 기준으로 잡으시는 게 현실적이에요.")
faq(doc, "20대에 윤곽수술을 하는 건 이른 편인가요?",
    "나이 자체보다는 골격 성장이 마무리됐는지, 보톡스 같은 비수술 방법으로 해결되는 범위인지가 기준이에요. "
    "검사 없이 나이만으로 판단하기는 어렵습니다.")
faq(doc, "시술을 계속 받고 있는데, 수술은 언제 고민해야 할까요?",
    "같은 시술을 반복해도 이전만큼 변화가 느껴지지 않는 시점이 하나의 신호가 될 수 있어요. "
    "다만 처짐 정도와 피부 상태에 따라 판단이 달라지니 진찰이 먼저입니다.")
faq(doc, "남자도 상안검이나 하안검 상담을 많이 받나요?",
    "50대 전후로 문의가 늘어나는 편이에요. 다만 남성은 눈꺼풀 두께와 피부 특성이 달라 접근이 같지 않으니, "
    "개별 진찰로 방향을 잡으셔야 합니다.")

heading(doc, "정리하면,")
para(doc, "나이대별 성형은 '몇 살에 뭘 해야 한다'의 문제가 아니라, "
          "'지금 내가 어디에서 불편함을 느끼고 있는가'의 문제입니다.")
para(doc, "10대와 20대는 눈, 20대 후반부터는 윤곽과 재수술, 40대 이후로는 눈꺼풀과 처짐. "
          "이 흐름은 참고용 지도일 뿐이고, 실제 순서는 내 얼굴을 정확히 본 뒤에 정해져요.")
para(doc, "지금 어디부터 시작해야 할지 고민이시라면, 혼자 순서를 정하려 하지 마시고 "
          "정확한 진단으로 나에게 맞는 우선순위를 확인해 보시길 바랍니다.")
para(doc, "무이성형외과에서는 특정 부위만 따로 보지 않고, 얼굴 전체 비율과 이전 수술 이력, "
          "지금의 피부 상태까지 함께 확인한 뒤 우선순위를 정리해 드립니다.")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(2)
_kf(p.add_run("함께 읽으면 좋은 글"), size=12, bold=True, color=GREEN)
bullet(doc, "무턱, 이중턱 짧은 턱끝, 턱끝성형으로 V라인 만드는 법 — https://blog.naver.com/mooips/224333520810")
bullet(doc, "광대축소 사각턱 처짐, 줄이는 법은? (AI 논문 결과 확인) — https://blog.naver.com/mooips/224354639518")
bullet(doc, "눈 처짐 수술, 나에게 맞는 방법 찾는 법 — [3월 글 URL 넣기]")
para(doc, "※ 이 글은 눈·윤곽·재수술을 모두 언급하는 '허브 글'이라, 기존 글 3개로 내부 링크를 걸기에 가장 좋은 자리예요.",
     size=9.5, color=MUT)

divider(doc)
para(doc, "#나이대별성형  #20대성형  #30대성형  #40대성형  #성형순서  #쌍꺼풀재수술  #상안검  #남자성형  #강남성형외과  #무이성형외과",
     size=10.5, color=ACCENT, after=10, line=1.6)
para(doc, "본 내용은 시술 및 수술 관련 정보를 이해하기 쉽게 전달하기 위한 정보 전달 목적으로 의료법 제56조 1항을 준수하여, "
          "무이성형외과의원에서 직접 작성했습니다. 모든 시/수술은 개인마다 차이 및 부작용이 발생할 수 있으니 "
          "반드시 담당 의료진과 상담 후 진행하시길 권장드립니다.", size=9.5, color=MUT, line=1.5)

divider(doc)

para(doc, "발행 체크리스트", size=12, bold=True, color=ACCENT, after=6)
table(doc, [
    ["✔", "항목", "세부"],
    ["☐", "주제 설정", "글쓰기 화면 우측 → '성형' 선택 (미설정 시 주제별 노출 누락)"],
    ["☐", "공개 설정", "전체공개 + 검색 허용 체크"],
    ["☐", "발행 시각", "오전 9~11시 또는 오후 8~10시"],
    ["☐", "영상 임베드", "글 상단(도입 아래)에 원본 유튜브 삽입 — 체류시간 상승"],
    ["☐", "표 재작성", "'나이대별 성형 지도' 표를 스마트에디터 표 기능으로 직접 작성"],
    ["☐", "캡처 이미지", "본문 캡처 6곳 삽입 + 각 이미지 ALT 입력"],
    ["☐", "내부 링크", "'함께 읽으면 좋은 글' 3건 연결 (눈 글 URL 확인)"],
    ["☐", "발행 후 금지", "제목·첫 문단 수정 금지 (재수집 시 순위 초기화 위험)"],
    ["☐", "D+1 확인", "'나이대별 성형 순서', '20대 30대 성형'으로 직접 검색해 노출 위치 기록"],
    ["☐", "D+3 확인", "노출 없으면 태그 1~2개만 교체 (본문·제목은 그대로)"],
], widths=[1.0, 3.2, 11.8])

para(doc, "", after=2)
para(doc, "v3 문체 교정 메모", size=12, bold=True, color=ACCENT, after=4, before=10)
bullet(doc, "인트로에 시그니처 복원: '오늘은 유일무이한 아름다움, 무이성형외과✨에서 ~ 솔직하게 풀어드릴게요' (107회 글과 동일 공식)")
bullet(doc, "'원장님이 ~ 풀어드릴게요' 3인칭 서술, '영상 내용 그대로' 같은 작업 지시문 흔적 전부 제거")
bullet(doc, "면책 문구를 인트로에서 빼고 표 안내('권장 순서표가 아니라는 점만')와 '다만' 섹션으로 자연스럽게 분산")
bullet(doc, "종결어미 리듬 교정: ~거든요 / ~예요 / ~하죠 / ~보시면 됩니다 (기준 글 224333520810 패턴)")
bullet(doc, "정리하면: '혼자 ~하지 마시고 → 무이성형외과에서는' 마무리 공식 복원")

out = os.path.abspath(os.path.join(os.path.dirname(__file__), "..",
        "무이성형외과_블로그초안_나이대별성형_리라이트.docx"))
doc.save(out)
print("saved:", out)
