# -*- coding: utf-8 -*-
"""
무이성형외과 유튜브 → 네이버 블로그 포스팅 초안 (윤곽수술 위험성·처짐 편)
소재 영상: "윤곽수술, 위험성과 처짐이 걱정 된다면?"
어투·양식 기준: 기존 무이성형외과 블로그 포스팅
이미지: 사용자 제공 5장(썸네일 / AI 논문 표지 / 랜드마크 / Fig.3 / Fig.4) → 자리표시로 배치
출력: ../무이성형외과_블로그초안_윤곽수술_위험성처짐.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

FONT = "Malgun Gothic"
INK = RGBColor(0x2B, 0x26, 0x22)
ACCENT = RGBColor(0xB0, 0x8A, 0x5E)
MUT = RGBColor(0x8A, 0x81, 0x78)
BODY = RGBColor(0x33, 0x30, 0x2C)
IMGBG = RGBColor(0x9A, 0x7A, 0x4E)

VIDEO_TITLE = "윤곽수술, 위험성과 처짐이 걱정 된다면?"


def _kfont(run, size=11, bold=False, color=BODY, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), FONT)
    return run


def para(doc, text="", size=11, bold=False, color=BODY, align=None,
         before=4, after=8, italic=False, line=1.5):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = line
    if text:
        _kfont(p.add_run(text), size=size, bold=bold, color=color, italic=italic)
    return p


def heading(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(20); pf.space_after = Pt(6); pf.line_spacing = 1.4
    _kfont(p.add_run(text), size=14, bold=True, color=INK)
    return p


def quote(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(2); pf.space_after = Pt(2); pf.line_spacing = 1.4
    pf.left_indent = Cm(0.4)
    _kfont(p.add_run(text), size=11.5, bold=True, color=INK)
    return p


def imgph(doc, label, caption=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(12); pf.space_after = Pt(3)
    _kfont(p.add_run(f"[ {label} ]"), size=10.5, bold=True, color=IMGBG)
    if caption:
        pc = doc.add_paragraph(); pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pc.paragraph_format.space_after = Pt(11)
        _kfont(pc.add_run(caption), size=9.5, color=MUT, italic=True)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.space_before = Pt(1); pf.space_after = Pt(3); pf.line_spacing = 1.4
    _kfont(p.add_run(text), size=11, color=BODY)


def faq(doc, q, a):
    pq = doc.add_paragraph()
    pq.paragraph_format.space_before = Pt(10); pq.paragraph_format.space_after = Pt(2)
    pq.paragraph_format.line_spacing = 1.4
    _kfont(pq.add_run("Q. "), size=11, bold=True, color=ACCENT)
    _kfont(pq.add_run(q), size=11, bold=True, color=INK)
    pa = doc.add_paragraph()
    pa.paragraph_format.space_after = Pt(4); pa.paragraph_format.line_spacing = 1.5
    _kfont(pa.add_run("A. "), size=11, bold=True, color=ACCENT)
    _kfont(pa.add_run(a), size=11, color=BODY)


def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pbdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pbdr.makeelement(qn('w:bottom'), {})
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'D9D2C8')
    pbdr.append(bottom); pPr.append(pbdr)


# ─────────────────────────────────────────────
doc = Document()
normal = doc.styles["Normal"]
normal.font.name = FONT; normal.font.size = Pt(11)
normal.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

# 상단 안내
para(doc, "무이성형외과 · 유튜브 영상 기반 네이버 블로그 포스팅 초안", size=10.5, color=MUT, after=2, before=0)
para(doc, f"소재 영상: {VIDEO_TITLE}   |   작성일: 2026-07-22", size=10.5, color=MUT, after=2, before=0)
para(doc, "※ [ ] 표시는 제공해 주신 원본 이미지를 넣을 위치입니다(①~⑤). 게시 전 의료광고 규정 최종 검수를 권장드립니다.",
     size=10.5, color=MUT, after=6, before=0)
divider(doc)

# 제목
para(doc, "윤곽수술 위험성과 처짐, 진짜 그럴까? AI 논문으로 직접 확인했습니다",
     size=17, bold=True, color=INK, before=2, after=6, line=1.3)
para(doc, "ㅣ광대축소·사각턱 등 안면윤곽수술을 고민 중이라면", size=11.5, color=MUT, after=12, before=0)

# 도입 인용 (썸네일 문구 그대로)
quote(doc, "“(위험성 있지 않아요?)”")
quote(doc, "“(처질까 봐 너무 걱정돼요)”")
quote(doc, "“(4, 50대도 윤곽수술 받나요?)”")
quote(doc, "“(저같이 고민 많이 하는 케이스 있었어요?)”")

para(doc, "윤곽수술 상담에서 가장 많이 나오는 이야기가 바로 이 두 가지예요. "
          "‘뼈를 건드리는 수술인데 위험하지 않을까’, 그리고 ‘수술하고 나면 얼굴이 처지지 않을까’ 하는 걱정이요. "
          "고민이 길어질 수밖에 없는 부분이라, 오늘은 이 두 가지를 하나씩 짚어드릴게요.", before=10)

para(doc, "결론부터 말씀드리면, 윤곽수술은 막연히 느껴지는 것만큼 무작정 위험한 수술은 아니고, "
          "처짐 역시 수술 방식과 개인 조건에 따라 충분히 줄여갈 수 있는 부분입니다.")

para(doc, "오늘은 유일무이한 아름다움, 무이성형외과✨에서 윤곽수술의 위험성과 처짐을 "
          "어떤 기준으로 바라보는지, 그리고 실제 연구 결과는 어땠는지 솔직하게 풀어드리겠습니다.")
imgph(doc, "이미지 ① 대표(썸네일)",
      "“윤곽수술, 위험성과 처짐이 걱정 된다면?” — 원장 정면 + 환자 고민 문구 4컷")

# 1. 위험성
heading(doc, "윤곽수술은 위험하다? — 왜 그렇게 느껴질까")
para(doc, "윤곽수술은 뼈를 다루는 수술이다 보니, 아무래도 ‘무섭다’, ‘위험하다’고 느끼기 쉬워요. "
          "하지만 실제로 윤곽수술은 전신마취를 하고 진행되며, 마취과 전문의가 호흡을 비롯한 전신 상태를 "
          "안정적으로 관리하는 가운데 이루어집니다.")
para(doc, "수술을 집도하는 입장에서 특히 신경 써서 보호하는 부분은 ‘신경선’이에요. "
          "신경선 외에 크게 조심해야 할 요소가 아주 많은 것은 아니어서, 신경을 잘 보호하며 진행한다면 "
          "위험 요소를 줄이면서 수술 시간도 길지 않게 진행할 수 있습니다. "
          "물론 모든 수술이 위험이 전혀 없다고 말할 수는 없기에, 정확한 진단과 상담이 먼저입니다.")

# 2. 처짐 + 논문
heading(doc, "윤곽수술 하면 처진다? — 직접 AI 논문으로 확인했어요")
para(doc, "윤곽수술을 하면 얼굴이 처진다고 걱정하시는 분들이 정말 많습니다. "
          "저희는 기본적으로 윤곽수술 자체가 얼굴을 처지게 한다고 보지는 않는데요. "
          "수술할 때 박리를 최소화하고, 얼굴 형태를 잡아주는 ‘유지인대’를 최대한 보존하며 진행하기 때문이에요. "
          "그래서 ‘처짐’이라기보다 ‘볼륨의 이동’에 가깝다고 설명드리는 편입니다.")
para(doc, "그럼에도 걱정하시는 분들이 많아, 실제로 AI를 활용해 수술 전후를 비교한 논문을 직접 작성해 확인해봤습니다. "
          "요즘 화두인 AI로 수술 전후 사진을 분석해, 정말 처졌는지 아닌지를 객관적으로 살펴본 연구예요. "
          "국제 학술지(PRS Global Open)에 게재된 내용입니다.")
imgph(doc, "이미지 ② AI 논문 표지",
      "PRS Global Open ‘A Comprehensive Assessment of Soft-tissue Sagging after Zygoma Reduction Surgery through AI Analysis’")
para(doc, "분석은 AI로 얼굴에 여러 기준점을 찍고, 네 가지 방법으로 진행했어요. "
          "팔자주름이 깊어지는지, 얼굴 아래쪽 굴곡이 늘어나는지 등을 수술 전후로 비교한 것이죠.")
imgph(doc, "이미지 ③ AI 분석 랜드마크",
      "수술 전 얼굴에 분석 기준점을 표시한 이미지")
para(doc, "결과적으로, 수술 전후에 통계적으로 유의미한 처짐 차이는 확인되지 않았습니다. "
          "아랫볼의 곡률, 팔자주름과 마리오네트 라인 깊이 모두 ‘의미 있는 처짐’으로 보긴 어렵다는 결론이었어요.")
imgph(doc, "이미지 ④ 논문 Fig.3",
      "아랫볼 곡률(Hessian filter) 분석 — 수술 전(A)·후(B) 비교")
imgph(doc, "이미지 ⑤ 논문 Fig.4",
      "팔자주름·마리오네트 라인(face mesh) 분석 — 수술 전(A)·후(B) 비교")

# 3. 처짐 두 경우
heading(doc, "그래도 처짐이 느껴지는 두 가지 경우 (나이 · 얼굴 지방)")
para(doc, "다만 수술을 많이 진행하다 보면, 처짐을 느끼시는 분들의 경향은 분명히 있습니다. 크게 두 가지예요.")
para(doc, "첫 번째는 ‘나이’입니다. 고무줄을 떠올리면 쉬운데요. 새 고무줄은 탁 놓으면 잘 돌아가지만, "
          "오래된 고무줄은 탄성이 떨어지죠. 20~30대는 탄력이 좋아 뼈를 줄여도 피부가 안쪽으로 자연스럽게 "
          "따라 들어가며 회복되는 편이지만, 나이가 조금 있으면 탄성이 떨어져 처짐을 느끼시는 경우가 있습니다.")
para(doc, "두 번째는 ‘얼굴 지방’이에요. 얼굴에 살이 많으면 아무래도 중력의 영향을 더 받기 때문에, "
          "처지는 경향이 조금 더 높아질 수 있습니다.")

# 4. 4050
heading(doc, "40·50대도 윤곽수술을 많이 받는 이유")
para(doc, "그럼에도 요즘은 40대, 50대 윤곽수술이 많이 늘고 있어요. "
          "울세라·써마지, 거상이나 실리프팅처럼 함께 활용할 수 있는 방법이 다양해진 것이 한 가지 이유입니다.")
para(doc, "또 하나는, 윤곽 없이 거상만 받으려 해도 광대·턱 같은 뼈 구조가 높으면 당겼을 때 그 윤곽이 오히려 "
          "도드라지는 경우가 있어요. 그래서 윤곽을 먼저 진행하거나, 윤곽과 거상을 함께 진행하기도 합니다. "
          "‘윤곽수술은 처져서 안 한다’는 이야기는 이제 예전 트렌드에 가깝습니다.")

# 5. 처짐 예방
heading(doc, "처짐을 예방하며 윤곽수술 하는 방법")
para(doc, "얼굴에 살이 많은 편이라면, 다이어트도 좋지만 지방흡입을 함께 하거나, 광대의 경우 볼지방(심부볼) 제거처럼 "
          "지방에 대한 조작을 같이 해드리면 처지는 경향을 줄이는 데 도움이 될 수 있어요.")
para(doc, "얼굴은 전체적으로 작아지는 방향이 유리하기 때문에, 얼굴 뼈도 줄이고 지방도 함께 줄여가는 방향으로 "
          "설계하면 만족도 측면에서도 도움이 되는 편입니다. 물론 이 조합은 개인의 상태에 따라 달라집니다.")

# 6. 결정/회복 (체험담 아님, 일반 정보)
heading(doc, "윤곽수술, 결정까지 오래 고민하게 되는 수술")
para(doc, "윤곽수술은 결정하기까지 오래 고민하게 되는 수술이에요. 하기로 마음먹기 전까지 계속 망설이게 되고, "
          "안 하기로 하면 오히려 마음이 편해지기도 하니까요. 실제로 몇 년을 고민하다 결정하시는 분들도 많습니다.")
para(doc, "회복 기간이 부담이 되는 것도 사실이에요. 다만 붓기는 수술 직후부터 한 달, 세 달, 6개월, 1년에 걸쳐 "
          "서서히 빠지면서 얼굴 라인이 점차 자리를 잡아갑니다. 변화가 큰 수술인 만큼, 충분한 정보와 안전을 먼저 "
          "확인하고 나에게 맞는 시점에 결정하시는 것을 권해드립니다.")

# 이런 분들
heading(doc, "이런 고민이라면 상담을 권해드려요")
bullet(doc, "윤곽수술의 위험성이 막연히 걱정되는 분")
bullet(doc, "수술 후 얼굴 처짐이 염려되는 분")
bullet(doc, "40·50대에 윤곽수술을 고민 중인 분")
bullet(doc, "얼굴 뼈와 지방을 함께 고려해 라인을 정리하고 싶은 분")

# FAQ
heading(doc, "자주 묻는 질문 (FAQ)")
faq(doc, "윤곽수술은 위험한가요?",
    "뼈를 다루는 수술이라 막연히 무섭게 느껴질 수 있어요. 전신마취 하에 마취과 전문의가 호흡 등 전신 상태를 관리하고, "
    "신경선을 보호하며 진행합니다. 다만 위험이 전혀 없는 수술은 없으므로, 정확한 진단과 충분한 상담이 먼저입니다.")
faq(doc, "윤곽수술을 하면 얼굴이 처지나요?",
    "박리를 최소화하고 유지인대를 보존하는 방향으로 진행합니다. AI로 수술 전후를 분석한 저희 논문에서도 통계적으로 "
    "유의미한 처짐 차이는 확인되지 않았어요. 다만 나이와 얼굴 지방량에 따라 개인차가 있을 수 있습니다.")
faq(doc, "4, 50대도 윤곽수술을 받을 수 있나요?",
    "최근에는 거상·리프팅 등과 병행하며 40·50대에도 많이 진행합니다. 뼈가 높은 경우 리프팅만으로는 한계가 있어, "
    "윤곽을 함께 고려하기도 해요.")
faq(doc, "처짐이 걱정되면 어떻게 하나요?",
    "얼굴 지방이 많은 경우 지방흡입이나 볼지방 제거 등을 함께 고려할 수 있어요. 나이·피부 탄력·지방량을 종합해 "
    "계획을 세우는 것이 중요합니다.")

# 정리
heading(doc, "정리하면,")
para(doc, "윤곽수술의 위험성과 처짐은 막연한 걱정이 큰 부분이에요. 전신마취 관리와 신경 보호 아래 진행되고, "
          "처짐도 최소 박리·유지인대 보존, 그리고 필요 시 지방 조작을 함께 고려해 관리해갈 수 있습니다.")
para(doc, "다만 나이·피부 탄력·얼굴 지방 같은 개인 조건에 따라 결과와 경과는 달라질 수 있어요. "
          "혼자 판단하기보다, 정확한 진단을 바탕으로 전문의와 충분히 상담하신 뒤 결정하시길 바랍니다.")

# 해시태그
para(doc, "#윤곽수술 #윤곽수술위험성 #윤곽수술처짐 #광대축소 #사각턱수술 #안면윤곽 #얼굴처짐 #40대윤곽 "
          "#50대윤곽 #거상수술 #실리프팅 #Vライン #강남윤곽수술 #강남성형외과 #무이성형외과".replace("Vライン", "V라인"),
     size=10.5, color=ACCENT, before=16, after=14, line=1.6)

divider(doc)
para(doc, "본 내용은 시술 및 수술 관련 정보를 이해하기 쉽게 전달하기 위한 정보 전달 목적으로 "
          "의료법 제56조 1항을 준수하여, 무이성형외과의원에서 직접 작성했습니다. "
          "모든 시/수술은 개인마다 차이 및 부작용이 발생할 수 있으니 반드시 담당 의료진과 상담 후 진행하시길 권장드립니다.",
     size=9.5, color=MUT, line=1.5)

out = os.path.abspath(os.path.join(os.path.dirname(__file__), "..",
        "무이성형외과_블로그초안_윤곽수술_위험성처짐.docx"))
doc.save(out)
print("saved:", out)
